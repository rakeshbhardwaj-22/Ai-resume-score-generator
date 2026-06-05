import re
import io
import os
import pypdf
import docx
from typing import Dict, Any, Union, List

class ResumeParser:
    """
    A class to parse resumes in PDF and DOCX formats.
    Extracts text, cleans it, identifies contact details, and partitions it into standard sections.
    """

    # Section keyword mappings for segmentation
    SECTION_HEADERS = {
        "skills": [
            r"\bskills\b", r"\btechnical skills\b", r"\bskills & technologies\b", 
            r"\bcore competencies\b", r"\btechnologies\b", r"\bexpertise\b", 
            r"\btools & technologies\b", r"\bprogramming languages\b"
        ],
        "experience": [
            r"\bexperience\b", r"\bwork experience\b", r"\bprofessional experience\b", 
            r"\bemployment history\b", r"\bwork history\b", r"\bcareer history\b"
        ],
        "education": [
            r"\beducation\b", r"\bacacademic background\b", r"\bacacademic credentials\b", 
            r"\bqualifications\b", r"\bacademics\b"
        ],
        "projects": [
            r"\bprojects\b", r"\bacacademic projects\b", r"\bkey projects\b", 
            r"\bpersonal projects\b", r"\bselected projects\b"
        ],
        "certifications": [
            r"\bcertifications\b", r"\bcertificates\b", r"\blicenses & certifications\b", 
            r"\baccreditations\b"
        ]
    }

    def __init__(self):
        pass

    def extract_text(self, file_source: Union[str, io.BytesIO, bytes], file_name: str = "") -> str:
        """
        Extract raw text from PDF or DOCX file.
        
        Args:
            file_source: File path (str), file-like object (BytesIO), or raw bytes.
            file_name: Optional file name to determine file type if file_source is bytes/stream.
            
        Returns:
            str: Extracted raw text.
        """
        # Determine the file type
        ext = ""
        if isinstance(file_source, str):
            _, ext = os.path.splitext(file_source.lower())
        elif file_name:
            _, ext = os.path.splitext(file_name.lower())
        else:
            # Fallback check if it's bytes
            raise ValueError("Unable to determine file type. Provide a file path or specify file_name.")

        # Ensure we have a file-like object if it's bytes
        if isinstance(file_source, bytes):
            file_source = io.BytesIO(file_source)

        if ext == ".pdf":
            return self._extract_from_pdf(file_source)
        elif ext in [".docx", ".doc"]:
            # Note: doc extension is treated as docx since python-docx only supports docx.
            # True .doc is binary and python-docx will throw an error, which we catch.
            return self._extract_from_docx(file_source)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")

    def _extract_from_pdf(self, file_source: Union[str, io.BytesIO]) -> str:
        """Helper to extract text from PDF using pypdf."""
        text = ""
        try:
            reader = pypdf.PdfReader(file_source)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {str(e)}")
        return text

    def _extract_from_docx(self, file_source: Union[str, io.BytesIO]) -> str:
        """Helper to extract text from DOCX using python-docx."""
        text = ""
        try:
            doc = docx.Document(file_source)
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {str(e)}")
        return text

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize raw extracted text.
        """
        if not text:
            return ""
        
        # Replace multiple spaces with single space
        cleaned = re.sub(r'[ \t]+', ' ', text)
        
        # Replace multiple newlines with a single newline or standard spacing
        cleaned = re.sub(r'\n+', '\n', cleaned)
        
        # Remove non-printable / control characters (except newline, carriage return, tab)
        cleaned = "".join(ch for ch in cleaned if ch.isprintable() or ch in ['\n', '\r', '\t'])
        
        return cleaned.strip()

    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """
        Extract Name, Email, and Phone number from resume text.
        """
        info = {"name": "Unknown", "email": "", "phone": ""}
        
        # 1. Extract Email
        email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
        emails = re.findall(email_pattern, text)
        if emails:
            info["email"] = emails[0].strip()

        # 2. Extract Phone
        # Standard phone patterns including country codes, spaces, dashes, parentheses
        phone_pattern = r'(\+?\d{1,4}[-.\s]?\(?\d{1,3}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9})'
        phones = re.findall(phone_pattern, text)
        for ph in phones:
            # Clean non-digit characters to check length
            cleaned_ph = re.sub(r'\D', '', ph)
            if 7 <= len(cleaned_ph) <= 15:
                info["phone"] = ph.strip()
                break

        # 3. Guess Name
        # In typical resumes, the candidate's name is on the first few lines of text.
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines[:5]:
            # Skip lines that contain email, phone, or are too long or short
            if "@" in line or any(d.isdigit() for d in line):
                continue
            if len(line.split()) <= 4 and len(line) > 2:
                # Basic check for title casing (e.g. John Doe)
                if all(part[0].isupper() for part in line.split() if part.isalpha()):
                    info["name"] = line
                    break

        return info

    def parse_sections(self, text: str) -> Dict[str, str]:
        """
        Segment resume text into logical sections based on predefined headers.
        
        Returns:
            Dict[str, str]: Sections dictionary (e.g., {"skills": "...", "experience": "..."})
        """
        lines = text.split('\n')
        section_indices = []
        
        # Normalize headers and find where they occur
        for line_num, line in enumerate(lines):
            clean_line = line.strip().lower()
            if not clean_line or len(clean_line) > 40:
                continue
            
            # Check if this line matches any section header pattern
            for section_name, patterns in self.SECTION_HEADERS.items():
                matched = False
                for pattern in patterns:
                    if re.match(pattern, clean_line):
                        section_indices.append((line_num, section_name))
                        matched = True
                        break
                if matched:
                    break

        # Sort found headings by line number
        section_indices.sort(key=lambda x: x[0])
        
        sections = {
            "header": "",
            "skills": "",
            "experience": "",
            "education": "",
            "projects": "",
            "certifications": ""
        }
        
        if not section_indices:
            # If no sections could be identified, put everything in the header section
            sections["header"] = text
            return sections
        
        # Extract content before the first section header
        first_idx = section_indices[0][0]
        sections["header"] = "\n".join(lines[:first_idx]).strip()
        
        # Extract content between section headers
        for i in range(len(section_indices)):
            curr_idx, curr_sec = section_indices[i]
            if i + 1 < len(section_indices):
                next_idx, _ = section_indices[i+1]
                section_text = "\n".join(lines[curr_idx + 1 : next_idx]).strip()
            else:
                section_text = "\n".join(lines[curr_idx + 1 :]).strip()
                
            # Append if section is repeated (e.g., multiple Skills headers)
            if sections[curr_sec]:
                sections[curr_sec] += "\n" + section_text
            else:
                sections[curr_sec] = section_text

        return sections

    def parse(self, file_source: Union[str, io.BytesIO, bytes], file_name: str = "") -> Dict[str, Any]:
        """
        Complete parsing pipeline for a resume file.
        
        Returns:
            Dict[str, Any]: Structured dictionary with name, email, phone, raw_text, cleaned_text, and sections.
        """
        raw_text = self.extract_text(file_source, file_name)
        cleaned_text = self.clean_text(raw_text)
        contact_info = self.extract_contact_info(cleaned_text)
        sections = self.parse_sections(cleaned_text)
        
        return {
            "name": contact_info["name"],
            "email": contact_info["email"],
            "phone": contact_info["phone"],
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "sections": sections
        }
